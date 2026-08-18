"""
Dựng lại biểu mẫu thành Markdown + DOCX + PDF của mình.

VÌ SAO DỰNG LẠI CHỨ KHÔNG ĐĂNG HTML GỐC. Ruột biểu mẫu là phụ lục của văn bản quy
phạm — Điều 15 Luật Sở hữu trí tuệ loại khỏi đối tượng bảo hộ, đúng căn cứ mà
README dùng cho trang công khai. Nhưng bản CHUYỂN ĐỔI sang HTML là công sức của
Thư viện Pháp luật. Nên bản đăng ra ngoài phải là bản dựng lại từ nội dung, kèm
ghi nguồn và link ngược; HTML gốc chỉ ở lại data/forms/html/ làm nguyên liệu.

VÌ SAO CÓ DOCX. Biểu mẫu sinh ra để ĐIỀN. PDF không điền được, nên nếu chỉ có PDF
thì người dùng vẫn phải sang TVPL tải bản Word — tính năng coi như không giải
quyết được việc gì.

VÌ SAO KHÔNG DÙNG build_report_pdf() CỦA src/utils/report_pdf.py. Bộ đó vẽ bìa
báo cáo với nhãn in cứng "BÁO CÁO PHÁP LÝ CHUYÊN ĐỀ THEO NGÀNH" và các ô Ngành /
Kỳ / Chốt số liệu. In nhãn đó lên một tờ hợp đồng khoán việc là nói sai bản chất
tài liệu. `convert_md_to_pdf()` của pdf_exporter nhận `report_title` tự do nên
hợp hơn.

BẢNG LÀ TOÀN BỘ NỘI DUNG CỦA BIỂU MẪU. Mất bảng là mất biểu mẫu, nên cả ba định
dạng đều đi qua cùng một bước dựng lưới có xử lý rowspan/colspan
(`luoi_bang()`) thay vì mỗi bên tự bóc một kiểu.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag

from src.config import FORMS_BUILD_DIR
from src.utils.pdf_exporter import convert_md_to_pdf

logger = logging.getLogger(__name__)

NGUON_TVPL = "Thư viện Pháp luật"

#: Bảng rộng hơn ngần này cột thì bảng pipe của Markdown không đọc nổi trên màn
#: hình. Vẫn dựng, nhưng DOCX mới là bản dùng được — ghi chú thẳng vào file.
COT_TOI_DA_DE_DOC = 12


@dataclass
class O:
    """Một ô sau khi đã trải lưới."""

    text: str = ""
    rowspan: int = 1
    colspan: int = 1
    #: True ở ô gốc, False ở các ô bị ô gốc phủ lên (do rowspan/colspan).
    goc: bool = True


@dataclass
class KetQuaDung:
    md_path: Path | None = None
    docx_path: Path | None = None
    pdf_path: Path | None = None
    canh_bao: list[str] = field(default_factory=list)


# ──────────────────────────────────────────────
# Bóc cấu trúc
# ──────────────────────────────────────────────
_THE_NGAT = ("br",)
#: Thẻ kết thúc một dòng. Phải ngắt sau chúng, KHÔNG chỉ ngắt ở <br>.
_THE_KHOI = ("p", "div", "tr", "li", "table", "h1", "h2", "h3", "h4", "h5", "h6")


def _chu(node: Tag | NavigableString) -> str:
    """Chữ của một nút: thẻ nội tuyến ghép sát, thẻ khối ngắt dòng.

    Ghép SÁT các thẻ nội tuyến vì Aspose cắt chữ thành nhiều `<span>`, có chỗ
    cắt giữa từ ("Độc lậ" + "p -"); chèn dấu cách vào đó là làm vỡ chữ.

    Nhưng phải NGẮT sau thẻ khối, nếu không hai `<p>` trong cùng một ô dính vào
    nhau: đã gặp thật ở biểu mẫu Kho bạc — "0124.N.KBNN" + "Kèm theo Thông tư…"
    ra thành "0124.N.KBNNKèm theo Thông tư…". Chữ dính kiểu này vừa khó đọc vừa
    làm bộ đếm từ khoá trượt.
    """
    if isinstance(node, NavigableString):
        return str(node)
    phan = []
    for con in node.children:
        if isinstance(con, NavigableString):
            phan.append(str(con))
        elif con.name in _THE_NGAT:
            phan.append("\n")
        else:
            phan.append(_chu(con))
            if con.name in _THE_KHOI:
                phan.append("\n")
    return "".join(phan)


def _gon(s: str) -> str:
    return re.sub(r"[ \t ]+", " ", (s or "").replace("\r", "")).strip()


def _o_mot_dong(node) -> str:
    """Chữ của một ô, gộp về MỘT dòng.

    Bảng pipe của Markdown vỡ hoàn toàn nếu ô chứa ký tự xuống dòng, và dấu "|"
    trong nội dung cũng phải thoát — nếu không một ô có "|" sẽ tách thành hai cột
    và cả bảng lệch từ đó trở đi.
    """
    return _gon(_chu(node).replace("\n", " ")).replace("|", "\\|")


def luoi_bang(bang: Tag) -> list[list[O]]:
    """Trải một `<table>` HTML thành lưới chữ nhật, giữ nguyên rowspan/colspan.

    Không trải lưới thì mọi hàng có ô gộp sẽ lệch cột, và biểu mẫu tài chính —
    loại dùng ô gộp dày đặc nhất — trở thành vô nghĩa. Ô bị phủ để RỖNG chứ không
    lặp lại nội dung ô gốc: lặp lại làm người đọc tưởng có nhiều giá trị.
    """
    luoi: list[list[O]] = []
    treo: dict[tuple[int, int], O] = {}   # (hàng, cột) → ô bị phủ do rowspan

    for i, tr in enumerate(bang.find_all("tr")):
        hang: list[O] = []
        cot = 0
        for td in tr.find_all(["td", "th"], recursive=False) or tr.find_all(["td", "th"]):
            while (i, cot) in treo:
                hang.append(treo.pop((i, cot)))
                cot += 1
            try:
                rs = max(1, int(td.get("rowspan", 1)))
                cs = max(1, int(td.get("colspan", 1)))
            except (TypeError, ValueError):
                rs = cs = 1
            o = O(text=_o_mot_dong(td), rowspan=rs, colspan=cs, goc=True)
            hang.append(o)
            for k in range(1, cs):
                hang.append(O(goc=False))
            for j in range(1, rs):
                for k in range(cs):
                    treo[(i + j, cot + k)] = O(goc=False)
            cot += cs
        while (i, cot) in treo:
            hang.append(treo.pop((i, cot)))
            cot += 1
        luoi.append(hang)

    rong = max((len(h) for h in luoi), default=0)
    for h in luoi:
        h.extend(O(goc=False) for _ in range(rong - len(h)))
    return luoi


def _la_bang_rong(luoi: list[list[O]]) -> bool:
    return not any(o.text for h in luoi for o in h)


# ──────────────────────────────────────────────
# Markdown
# ──────────────────────────────────────────────
def _bang_sang_markdown(luoi: list[list[O]]) -> list[str]:
    if not luoi or _la_bang_rong(luoi):
        return []
    rong = len(luoi[0])
    dong = ["| " + " | ".join(o.text or " " for o in luoi[0]) + " |",
            "|" + "|".join(["---"] * rong) + "|"]
    for hang in luoi[1:]:
        dong.append("| " + " | ".join(o.text or " " for o in hang) + " |")
    return dong


def html_sang_markdown(body_html: str) -> str:
    """Ruột biểu mẫu → Markdown, bảng thành bảng pipe.

    Bảng pipe KHÔNG diễn tả được ô gộp — đó là hạn chế của chính Markdown, không
    phải của bộ chuyển này. Ô gộp vì vậy hiện thành ô rỗng bên cạnh ô có chữ.
    Bản DOCX giữ được ô gộp thật, nên nó mới là bản để điền.
    """
    soup = BeautifulSoup(body_html or "", "html.parser")
    ra: list[str] = []

    for bang in soup.find_all("table"):
        bang["data-da-xu-ly"] = "1"

    def duyet(node: Tag) -> None:
        for con in node.children:
            if isinstance(con, NavigableString):
                t = _gon(str(con))
                if t:
                    ra.append(t)
                continue
            if con.name == "table":
                ra.append("")
                ra.extend(_bang_sang_markdown(luoi_bang(con)))
                ra.append("")
                continue
            if con.find("table"):
                duyet(con)
                continue
            if con.name in ("p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6"):
                for dong in _chu(con).split("\n"):
                    t = _gon(dong)
                    if t:
                        ra.append(t)
                        ra.append("")
                continue
            duyet(con)

    duyet(soup)

    # Gộp dòng trống liên tiếp: Word xuất rất nhiều <p> rỗng làm giãn cách.
    sach: list[str] = []
    for dong in ra:
        if not dong and (not sach or not sach[-1]):
            continue
        sach.append(dong)
    return "\n".join(sach).strip()


def khoi_ghi_nguon(title: str, url: str, can_cu: list[str],
                   cap_nhat: str = "") -> str:
    """Khối nguồn bắt buộc có trên MỌI bản dựng lại.

    Đây là phần giữ cho việc dựng lại nằm đúng phía ranh giới: nội dung biểu mẫu
    là phụ lục văn bản quy phạm (không thuộc đối tượng bảo hộ), còn công chuyển
    đổi là của TVPL nên phải ghi nguồn và trỏ ngược về.
    """
    dong = [f"# {title}", ""]
    if can_cu:
        dong.append(f"**Căn cứ:** {', '.join(can_cu)}")
    if cap_nhat:
        dong.append(f"**Cập nhật:** {cap_nhat}")
    dong += [
        f"**Nguồn:** {NGUON_TVPL} — {url}",
        "",
        "> Bản dựng lại từ nội dung biểu mẫu để tiện điền và in. Nội dung biểu "
        "mẫu là phụ lục của văn bản quy phạm pháp luật. Khi có khác biệt, lấy "
        "bản tại văn bản gốc làm chuẩn.",
        "",
    ]
    return "\n".join(dong)


# ──────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────
def dung_docx(tieu_de: str, body_html: str, out_path: Path,
              phan_dau: str = "") -> Path:
    """Dựng .docx giữ nguyên bảng và ô gộp — bản để NGƯỜI DÙNG ĐIỀN."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    # Times New Roman 13pt: đúng thể thức văn bản hành chính Việt Nam, cũng là
    # font mà chính biểu mẫu gốc dùng.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    doc.add_heading(tieu_de, level=1)
    for dong in (phan_dau or "").splitlines():
        if dong.strip():
            doc.add_paragraph(dong.strip())

    soup = BeautifulSoup(body_html or "", "html.parser")

    def them_bang(bang_tag: Tag) -> None:
        luoi = luoi_bang(bang_tag)
        if not luoi or _la_bang_rong(luoi):
            return
        t = doc.add_table(rows=len(luoi), cols=len(luoi[0]))
        t.style = "Table Grid"
        for i, hang in enumerate(luoi):
            for j, o in enumerate(hang):
                if o.goc:
                    t.cell(i, j).text = o.text
        # Gộp ô SAU khi đã điền chữ: gộp trước thì các ô con biến mất khỏi lưới
        # và chỉ số (i, j) không còn trỏ đúng chỗ.
        for i, hang in enumerate(luoi):
            for j, o in enumerate(hang):
                if not o.goc or (o.rowspan == 1 and o.colspan == 1):
                    continue
                i2 = min(i + o.rowspan - 1, len(luoi) - 1)
                j2 = min(j + o.colspan - 1, len(hang) - 1)
                try:
                    t.cell(i, j).merge(t.cell(i2, j2))
                except Exception as e:      # lưới lệch do HTML hỏng
                    logger.debug("Không gộp được ô (%d,%d): %s", i, j, e)

    def duyet(node: Tag) -> None:
        for con in node.children:
            if isinstance(con, NavigableString):
                t = _gon(str(con))
                if t:
                    doc.add_paragraph(t)
                continue
            if con.name == "table":
                them_bang(con)
                continue
            if con.find("table"):
                duyet(con)
                continue
            if con.name in ("p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6"):
                for dong in _chu(con).split("\n"):
                    t = _gon(dong)
                    if t:
                        doc.add_paragraph(t)
                continue
            duyet(con)

    duyet(soup)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ──────────────────────────────────────────────
# Dựng cả bộ
# ──────────────────────────────────────────────
def thu_muc_dung(form_key: str) -> Path:
    return FORMS_BUILD_DIR / form_key


def dung_tat_ca(form_key: str, tieu_de: str, body_html: str, url: str,
                can_cu: list[str] | None = None, cap_nhat: str = "",
                lam_pdf: bool = True) -> KetQuaDung:
    """Dựng Markdown + DOCX + PDF cho một biểu mẫu.

    PDF hỏng KHÔNG làm hỏng cả mẻ: nó là bản phụ, còn DOCX mới là bản người dùng
    cần. Ghi cảnh báo rồi đi tiếp thay vì dừng cả lượt chạy hàng trăm mẫu.
    """
    kq = KetQuaDung()
    thu_muc = thu_muc_dung(form_key)
    thu_muc.mkdir(parents=True, exist_ok=True)

    dau = khoi_ghi_nguon(tieu_de, url, can_cu or [], cap_nhat)
    than = html_sang_markdown(body_html)

    kq.md_path = thu_muc / f"{form_key}.md"
    kq.md_path.write_text(dau + "\n" + than + "\n", encoding="utf-8")

    try:
        kq.docx_path = dung_docx(tieu_de, body_html, thu_muc / f"{form_key}.docx",
                                 phan_dau=dau.replace("#", "").strip())
    except Exception as e:
        kq.canh_bao.append(f"DOCX hỏng: {e}")
        logger.warning("%s: dựng DOCX hỏng — %s", form_key, e)

    if lam_pdf:
        try:
            kq.pdf_path = convert_md_to_pdf(
                dau + "\n" + than,
                output_path=thu_muc / f"{form_key}.pdf",
                report_title=tieu_de[:120],
            )
        except Exception as e:
            kq.canh_bao.append(f"PDF hỏng: {e}")
            logger.warning("%s: dựng PDF hỏng — %s", form_key, e)

    return kq


__all__ = [
    "NGUON_TVPL", "COT_TOI_DA_DE_DOC", "O", "KetQuaDung",
    "luoi_bang", "html_sang_markdown", "khoi_ghi_nguon",
    "dung_docx", "dung_tat_ca", "thu_muc_dung",
]
