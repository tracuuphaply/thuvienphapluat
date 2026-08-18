"""Dựng lại biểu mẫu thành Markdown + DOCX + PDF.

BẢNG LÀ TOÀN BỘ NỘI DUNG CỦA BIỂU MẪU — mất bảng là mất biểu mẫu. Phần lớn test ở
đây canh đúng chỗ đó: trải lưới rowspan/colspan, không lệch cột, không dính chữ.
"""
import gzip
from pathlib import Path

import pytest
from bs4 import BeautifulSoup

from src.forms.renderer import (
    NGUON_TVPL,
    dung_docx,
    dung_tat_ca,
    html_sang_markdown,
    khoi_ghi_nguon,
    luoi_bang,
)
from src.sources.tvpl_forms_parse import tach_chi_tiet

FIXTURES = Path(__file__).parent / "fixtures" / "forms"


def doc(ten: str) -> str:
    return gzip.decompress((FIXTURES / f"{ten}.html.gz").read_bytes()).decode("utf-8")


def bang(html: str):
    return luoi_bang(BeautifulSoup(html, "html.parser").find("table"))


class TestTraiLuoiBang:
    def test_bang_don_gian(self):
        luoi = bang("<table><tr><td>A</td><td>B</td></tr>"
                    "<tr><td>C</td><td>D</td></tr></table>")
        assert [[o.text for o in h] for h in luoi] == [["A", "B"], ["C", "D"]]

    def test_colspan_khong_lam_lech_cot(self):
        """Không trải lưới thì hàng có ô gộp ngắn hơn hàng khác và lệch từ đó đi."""
        luoi = bang("<table><tr><td colspan='2'>Gộp ngang</td></tr>"
                    "<tr><td>A</td><td>B</td></tr></table>")
        assert len(luoi[0]) == len(luoi[1]) == 2
        assert luoi[0][0].text == "Gộp ngang" and luoi[0][0].colspan == 2
        assert luoi[0][1].text == "" and not luoi[0][1].goc

    def test_rowspan_day_o_hang_sau_sang_phai(self):
        """Ô rowspan chiếm chỗ ở hàng dưới; bỏ qua là mọi ô hàng dưới lùi một cột."""
        luoi = bang("<table>"
                    "<tr><td rowspan='2'>Dọc</td><td>B1</td></tr>"
                    "<tr><td>B2</td></tr></table>")
        assert [[o.text for o in h] for h in luoi] == [["Dọc", "B1"], ["", "B2"]]

    def test_o_bi_phu_de_rong_khong_lap_lai_noi_dung(self):
        """Lặp lại nội dung ô gốc làm người đọc tưởng có nhiều giá trị khác nhau."""
        luoi = bang("<table><tr><td colspan='3'>X</td></tr></table>")
        assert [o.text for o in luoi[0]] == ["X", "", ""]

    def test_moi_hang_deu_bang_nhau_du_html_thieu_o(self):
        luoi = bang("<table><tr><td>A</td><td>B</td><td>C</td></tr>"
                    "<tr><td>D</td></tr></table>")
        assert len(luoi[0]) == len(luoi[1]) == 3

    def test_rowspan_la_bo_qua_khi_gia_tri_hong(self):
        luoi = bang("<table><tr><td rowspan='abc'>A</td><td>B</td></tr></table>")
        assert luoi[0][0].rowspan == 1


class TestGhepChuKhongLamVoTu:
    def test_khong_chen_dau_cach_vao_giua_tu(self):
        """Aspose cắt chữ giữa từ: <span>Độc lậ</span><span>p -</span>."""
        md = html_sang_markdown("<p><span>Độc lậ</span><span>p - Tự do</span></p>")
        assert "Độc lập - Tự do" in md

    def test_hai_doan_trong_cung_mot_o_khong_dinh_nhau(self):
        """LỖI ĐÃ XẢY RA THẬT trên biểu mẫu Kho bạc.

        Ngắt chỉ ở <br> thì "0124.N.KBNN" và "Kèm theo Thông tư…" dính thành
        "0124.N.KBNNKèm theo Thông tư…".
        """
        md = html_sang_markdown(
            "<table><tr><td><p>0124.N.KBNN</p><p>Kèm theo Thông tư</p></td></tr></table>"
        )
        assert "KBNNKèm" not in md
        assert "0124.N.KBNN Kèm theo" in md

    def test_o_co_dau_gach_dung_duoc_thoat(self):
        """Một dấu "|" trong nội dung sẽ tách ô thành hai và lệch cả bảng."""
        import re

        md = html_sang_markdown("<table><tr><td>A|B</td><td>C</td></tr></table>")
        dong_bang = [d for d in md.splitlines() if d.startswith("|")][0]
        # Đếm vách thật, tức dấu "|" KHÔNG bị thoát: 2 ô → 3 vách.
        assert len(re.findall(r"(?<!\\)\|", dong_bang)) == 3
        assert r"A\|B" in dong_bang

    def test_xuong_dong_trong_o_bi_gop_ve_mot_dong(self):
        """Bảng pipe vỡ hoàn toàn nếu ô chứa ký tự xuống dòng."""
        md = html_sang_markdown("<table><tr><td>A<br>B</td><td>C</td></tr></table>")
        assert len([d for d in md.splitlines() if d.startswith("|")]) == 2


class TestMarkdownTuBieuMauThat:
    def test_giu_du_bang_cua_mau_kho_bac(self):
        d = tach_chi_tiet(doc("bieumau_detail_47156_khobac"), "bieumau", "47156")
        so_bang = len(BeautifulSoup(d.body_html, "html.parser").find_all("table"))
        md = html_sang_markdown(d.body_html)
        so_bang_md = md.count("\n|---")
        assert so_bang >= 1
        assert so_bang_md == so_bang, "mất bảng là mất biểu mẫu"

    def test_giu_noi_dung_hop_dong(self):
        d = tach_chi_tiet(doc("hopdong_detail_46696_khoanviec"), "hopdong", "46696")
        md = html_sang_markdown(d.body_html)
        assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in md
        assert "Độc lập - Tự do - Hạnh phúc" in md
        assert "BÊN A" in md

    def test_khong_de_lai_dong_trong_lien_tiep(self):
        """Word xuất rất nhiều <p> rỗng; giữ nguyên thì file giãn gấp đôi."""
        d = tach_chi_tiet(doc("hopdong_detail_46696_khoanviec"), "hopdong", "46696")
        md = html_sang_markdown(d.body_html)
        assert "\n\n\n" not in md


class TestKhoiGhiNguon:
    def test_luon_co_nguon_va_link_nguoc(self):
        """Chốt ranh giới bản quyền: bản dựng lại phải ghi nguồn và trỏ ngược."""
        k = khoi_ghi_nguon("MẪU X", "https://thuvienphapluat.vn/bieumau/1/x",
                           ["131/2025/TT-BTC"], "18/08/2026")
        assert NGUON_TVPL in k
        assert "https://thuvienphapluat.vn/bieumau/1/x" in k
        assert "131/2025/TT-BTC" in k
        assert "văn bản gốc làm chuẩn" in k

    def test_khong_co_can_cu_van_dung_duoc(self):
        assert "Căn cứ" not in khoi_ghi_nguon("MẪU X", "https://x", [], "")


class TestDungFile:
    def test_docx_giu_bang_va_o_gop(self, tmp_path):
        from docx import Document as DocxDocument

        p = dung_docx("MẪU THỬ",
                      "<table><tr><td colspan='2'>Tiêu đề gộp</td></tr>"
                      "<tr><td>A</td><td>B</td></tr></table>",
                      tmp_path / "t.docx")
        d = DocxDocument(str(p))
        assert len(d.tables) == 1
        t = d.tables[0]
        assert t.cell(0, 0).text == "Tiêu đề gộp"
        # Ô gộp: hai chỉ số khác nhau trỏ về cùng một ô
        assert t.cell(0, 0)._tc is t.cell(0, 1)._tc

    def test_dung_ca_bo_tu_bieu_mau_that(self, tmp_path, monkeypatch):
        import src.forms.renderer as r

        monkeypatch.setattr(r, "FORMS_BUILD_DIR", tmp_path)
        d = tach_chi_tiet(doc("hopdong_detail_46696_khoanviec"), "hopdong", "46696")
        kq = r.dung_tat_ca("hopdong-46696", d.title, d.body_html,
                           "https://thuvienphapluat.vn/hopdong/46696/x",
                           can_cu=["91/2015/QH13"])
        assert kq.md_path.exists() and kq.md_path.stat().st_size > 1000
        assert kq.docx_path.exists() and kq.docx_path.stat().st_size > 5000
        assert kq.pdf_path.exists() and kq.pdf_path.stat().st_size > 5000
        assert kq.canh_bao == []

    def test_pdf_hong_khong_lam_hong_ca_me(self, tmp_path, monkeypatch):
        """PDF là bản phụ; DOCX mới là bản người dùng cần để điền.

        Dừng cả lượt chạy hàng trăm mẫu vì một PDF hỏng là mất cả mẻ.
        """
        import src.forms.renderer as r

        monkeypatch.setattr(r, "FORMS_BUILD_DIR", tmp_path)
        monkeypatch.setattr(r, "convert_md_to_pdf",
                            lambda *a, **kw: (_ for _ in ()).throw(RuntimeError("font")))
        kq = r.dung_tat_ca("x-1", "MẪU X", "<p>nội dung dài" + "x" * 300 + "</p>",
                           "https://x/")
        assert kq.docx_path and kq.docx_path.exists()
        assert kq.pdf_path is None
        assert any("PDF hỏng" in c for c in kq.canh_bao)


@pytest.mark.parametrize("html", ["", "<div></div>", "<table></table>"])
def test_ruot_rong_khong_lam_vo(html):
    assert isinstance(html_sang_markdown(html), str)
